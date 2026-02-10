"""
Resume Routes - Resume component management (experiences, bullets, skills, education)
"""
import json
from datetime import datetime
from html import escape as html_escape
from io import StringIO, BytesIO
from flask import Blueprint, render_template, request, redirect, url_for, flash, session, send_file
from models import Experience, Bullet, Skill, Education, Suggestion, ExportProfile
from models.database import get_db_context
from services.export_transform import apply_export_rules
from services import get_ai_service
from utils import save_uploaded_file, extract_text_from_file, cleanup_file
import markdown
from docx import Document
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

bp = Blueprint('resume', __name__, url_prefix='/resume')


# ============================================================================
# MAIN RESUME VIEW
# ============================================================================

@bp.route('/')
def view_resume():
    """View all resume components"""
    experiences = Experience.get_all()
    bullets = Bullet.get_all()
    skills = Skill.get_all()
    education = Education.get_all()
    
    return render_template(
        'resume.html',
        experiences=experiences,
        bullets=bullets,
        skills=skills,
        education=education
    )


# ============================================================================
# RESUME IMPORT
# ============================================================================

@bp.route('/import', methods=['GET', 'POST'])
def import_resume():
    """Import resume from file"""
    if request.method == 'POST':
        if 'resume_file' not in request.files:
            flash('No file uploaded', 'error')
            return redirect(request.url)
        
        file = request.files['resume_file']
        
        if file.filename == '':
            flash('No file selected', 'error')
            return redirect(request.url)
        
        try:
            # Save file
            filepath = save_uploaded_file(file)
            flash(f'File uploaded: {file.filename}', 'info')
            
            # Extract text
            flash('Extracting text from file...', 'info')
            resume_text = extract_text_from_file(filepath)
            flash(f'Extracted {len(resume_text)} characters', 'info')
            
            # Parse with AI
            flash('Analyzing resume with Claude AI...', 'info')
            ai = get_ai_service()
            parsed_data = ai.parse_resume(resume_text)
            
            if not parsed_data:
                flash('Could not parse resume. Please try again.', 'error')
                cleanup_file(filepath)
                return redirect(request.url)
            
            flash('Resume parsed successfully!', 'success')
            
            # Get AI suggestions
            flash('Getting AI suggestions...', 'info')
            suggestions = ai.get_resume_suggestions(parsed_data)
            
            if suggestions:
                flash('AI suggestions generated!', 'success')
            else:
                flash('Could not generate suggestions, but resume was parsed.', 'warning')
            
            # Store in session for review page
            session['parsed_resume'] = json.dumps(parsed_data)
            session['resume_suggestions'] = json.dumps(suggestions) if suggestions else None
            
            # Clean up uploaded file
            cleanup_file(filepath)
            
            return redirect(url_for('resume.review_import'))
            
        except ValueError as e:
            flash(str(e), 'error')
            return redirect(request.url)
        except Exception as e:
            flash(f'Error processing resume: {str(e)}', 'error')
            return redirect(request.url)
    
    return render_template('import_resume.html')


@bp.route('/import/review')
def review_import():
    """Review parsed resume before saving"""
    if 'parsed_resume' not in session:
        flash('No resume data found. Please upload a resume first.', 'error')
        return redirect(url_for('resume.import_resume'))
    
    parsed_data = json.loads(session['parsed_resume'])
    suggestions = (
        json.loads(session['resume_suggestions']) 
        if session.get('resume_suggestions') 
        else None
    )
    
    return render_template(
        'review_resume_import.html',
        data=parsed_data,
        suggestions=suggestions
    )


def check_experience_duplicate(company_name, job_title):
    """Check if an experience with the same company and title already exists"""
    with get_db_context() as (conn, cursor):
        cursor.execute(
            "SELECT id FROM experiences WHERE LOWER(company_name) = LOWER(?) AND LOWER(job_title) = LOWER(?)",
            (company_name, job_title)
        )
        return cursor.fetchone() is not None


def check_bullet_duplicate(bullet_text):
    """Check if a bullet with the same text already exists"""
    with get_db_context() as (conn, cursor):
        cursor.execute(
            "SELECT id FROM bullets WHERE LOWER(bullet_text) = LOWER(?)",
            (bullet_text,)
        )
        return cursor.fetchone() is not None


def check_skill_duplicate(skill_name):
    """Check if a skill with the same name already exists"""
    with get_db_context() as (conn, cursor):
        cursor.execute(
            "SELECT id FROM skills WHERE LOWER(skill_name) = LOWER(?)",
            (skill_name,)
        )
        return cursor.fetchone() is not None


def check_education_duplicate(school_name, degree, field_of_study):
    """Check if an education entry with the same details already exists"""
    with get_db_context() as (conn, cursor):
        cursor.execute(
            "SELECT id FROM education WHERE LOWER(school_name) = LOWER(?) AND LOWER(degree) = LOWER(?) AND LOWER(field_of_study) = LOWER(?)",
            (school_name, degree, field_of_study)
        )
        return cursor.fetchone() is not None


@bp.route('/import/save', methods=['POST'])
def save_import():
    """Save the reviewed resume components to database"""
    if 'parsed_resume' not in session:
        flash('No resume data found', 'error')
        return redirect(url_for('resume.import_resume'))

    parsed_data = json.loads(session['parsed_resume'])
    suggestions_data = (
        json.loads(session['resume_suggestions'])
        if session.get('resume_suggestions')
        else None
    )
    
    try:
        experience_map = {}
        bullet_map = {}

        # Track what was added vs skipped
        stats = {
            'experiences_added': 0,
            'experiences_skipped': 0,
            'bullets_added': 0,
            'bullets_skipped': 0,
            'skills_added': 0,
            'skills_skipped': 0,
            'education_added': 0,
            'education_skipped': 0
        }

        # Save experiences (skip duplicates)
        for exp in parsed_data.get('experiences', []):
            if check_experience_duplicate(exp['company'], exp['title']):
                stats['experiences_skipped'] += 1
                # Still map for bullet relationships, but get existing ID
                with get_db_context() as (conn, cursor):
                    cursor.execute(
                        "SELECT id FROM experiences WHERE LOWER(company_name) = LOWER(?) AND LOWER(job_title) = LOWER(?)",
                        (exp['company'], exp['title'])
                    )
                    result = cursor.fetchone()
                    if result:
                        experience_map[exp['company']] = result['id']
            else:
                exp_id = Experience.create(
                    company_name=exp['company'],
                    job_title=exp['title'],
                    start_date=exp.get('start_date', ''),
                    end_date=exp.get('end_date', ''),
                    location=exp.get('location', ''),
                    description=exp.get('description', '')
                )
                experience_map[exp['company']] = exp_id
                stats['experiences_added'] += 1

        # Save bullets (skip duplicates)
        for bullet in parsed_data.get('bullets', []):
            if check_bullet_duplicate(bullet['text']):
                stats['bullets_skipped'] += 1
                # Still map for suggestions
                with get_db_context() as (conn, cursor):
                    cursor.execute(
                        "SELECT id FROM bullets WHERE LOWER(bullet_text) = LOWER(?)",
                        (bullet['text'],)
                    )
                    result = cursor.fetchone()
                    if result:
                        bullet_map[bullet['text']] = result['id']
            else:
                exp_id = experience_map.get(bullet.get('experience_company'))
                bullet_id = Bullet.create(
                    bullet_text=bullet['text'],
                    experience_id=exp_id,
                    tags=bullet.get('tags', ''),
                    category=bullet.get('category', '')
                )
                bullet_map[bullet['text']] = bullet_id
                stats['bullets_added'] += 1

        # Save skills (skip duplicates)
        for skill in parsed_data.get('skills', []):
            if check_skill_duplicate(skill['name']):
                stats['skills_skipped'] += 1
            else:
                Skill.create(
                    skill_name=skill['name'],
                    category=skill.get('category', '')
                )
                stats['skills_added'] += 1

        # Save education (skip duplicates)
        for edu in parsed_data.get('education', []):
            if check_education_duplicate(edu['school'], edu.get('degree', ''), edu.get('field', '')):
                stats['education_skipped'] += 1
            else:
                Education.create(
                    school_name=edu['school'],
                    degree=edu.get('degree', ''),
                    field_of_study=edu.get('field', ''),
                    graduation_year=edu.get('graduation_year', ''),
                    location=edu.get('location', '')
                )
                stats['education_added'] += 1
        
        # Save AI suggestions
        suggestion_count = 0
        if suggestions_data:
            # Experience alternate titles
            for exp_sugg in suggestions_data.get('experience_suggestions', []):
                exp_id = experience_map.get(exp_sugg['company'])
                if exp_id and exp_sugg.get('alternate_titles'):
                    for alt_title in exp_sugg['alternate_titles']:
                        Suggestion.create(
                            suggestion_type=Suggestion.TYPE_EXPERIENCE_ALT_TITLES,
                            component_id=exp_id,
                            original_text=exp_sugg['current_title'],
                            suggested_text=alt_title.strip(),
                            reasoning='AI-suggested alternate title to appeal to different roles'
                        )
                        suggestion_count += 1
            
            # Bullet improvements
            for bullet_sugg in suggestions_data.get('bullet_suggestions', []):
                bullet_id = bullet_map.get(bullet_sugg['original'])
                if bullet_id:
                    Suggestion.create(
                        suggestion_type=Suggestion.TYPE_BULLET_IMPROVEMENT,
                        component_id=bullet_id,
                        original_text=bullet_sugg['original'],
                        suggested_text=bullet_sugg['improved'],
                        reasoning=bullet_sugg['reason']
                    )
                    suggestion_count += 1
            
            # New skills
            for skill_sugg in suggestions_data.get('skill_suggestions', []):
                Suggestion.create(
                    suggestion_type=Suggestion.TYPE_NEW_SKILL,
                    suggested_text=skill_sugg,
                    reasoning='AI-suggested skill based on your experience'
                )
                suggestion_count += 1
            
            # Clarifying questions
            for question in suggestions_data.get('clarifying_questions', []):
                Suggestion.create(
                    suggestion_type=Suggestion.TYPE_CLARIFYING_QUESTION,
                    suggested_text=question,
                    reasoning='Question to help improve your profile'
                )
                suggestion_count += 1
        
        # Clear session
        session.pop('parsed_resume', None)
        session.pop('resume_suggestions', None)

        # Build success message with stats
        total_added = (stats['experiences_added'] + stats['bullets_added'] +
                      stats['skills_added'] + stats['education_added'])
        total_skipped = (stats['experiences_skipped'] + stats['bullets_skipped'] +
                        stats['skills_skipped'] + stats['education_skipped'])

        success_msg = f'Resume imported! Added {total_added} new component(s)'
        if total_skipped > 0:
            success_msg += f', skipped {total_skipped} duplicate(s)'
        success_msg += f'. {suggestion_count} AI suggestions saved.'

        # Add detailed breakdown if there were skips
        if total_skipped > 0:
            details = []
            if stats['experiences_skipped'] > 0:
                details.append(f"{stats['experiences_skipped']} experience(s)")
            if stats['bullets_skipped'] > 0:
                details.append(f"{stats['bullets_skipped']} bullet(s)")
            if stats['skills_skipped'] > 0:
                details.append(f"{stats['skills_skipped']} skill(s)")
            if stats['education_skipped'] > 0:
                details.append(f"{stats['education_skipped']} education entry(s)")

            if details:
                flash(f'Skipped duplicates: {", ".join(details)}', 'info')

        flash(success_msg, 'success')
        return redirect(url_for('resume.view_resume'))
        
    except Exception as e:
        flash(f'Error saving resume: {str(e)}', 'error')
        return redirect(url_for('resume.review_import'))


# ============================================================================
# EXPERIENCE CRUD
# ============================================================================

@bp.route('/experience/add', methods=['GET', 'POST'])
def add_experience():
    """Add a new experience"""
    if request.method == 'POST':
        Experience.create(
            company_name=request.form['company'],
            job_title=request.form['title'],
            alternate_titles=request.form.get('alt_titles', ''),
            start_date=request.form.get('start_date', ''),
            end_date=request.form.get('end_date', ''),
            location=request.form.get('location', ''),
            description=request.form.get('description', '')
        )
        flash('Experience added successfully', 'success')
        return redirect(url_for('resume.view_resume'))
    
    return render_template('add_experience.html')


@bp.route('/experience/<int:exp_id>/edit', methods=['GET', 'POST'])
def edit_experience(exp_id):
    """Edit an experience"""
    if request.method == 'POST':
        if Experience.update(
            exp_id,
            company_name=request.form['company'],
            job_title=request.form['title'],
            alternate_titles=request.form.get('alt_titles', ''),
            start_date=request.form.get('start_date', ''),
            end_date=request.form.get('end_date', ''),
            location=request.form.get('location', ''),
            description=request.form.get('description', '')
        ):
            flash('Experience updated successfully', 'success')
        else:
            flash('Experience not found', 'error')
        return redirect(url_for('resume.view_resume'))
    
    experience = Experience.get_by_id(exp_id)
    if not experience:
        flash('Experience not found', 'error')
        return redirect(url_for('resume.view_resume'))
    
    return render_template('edit_experience.html', exp=experience)


@bp.route('/experience/<int:exp_id>/delete', methods=['POST'])
def delete_experience(exp_id):
    """Delete an experience"""
    if Experience.delete(exp_id):
        flash('Experience deleted successfully', 'success')
    else:
        flash('Experience not found', 'error')
    return redirect(url_for('resume.view_resume') + '#experiences-section')


# ============================================================================
# BULLET CRUD
# ============================================================================

@bp.route('/bullet/add', methods=['GET', 'POST'])
def add_bullet():
    """Add a new bullet"""
    if request.method == 'POST':
        exp_id = request.form.get('experience_id')
        if exp_id == '':
            exp_id = None
        
        Bullet.create(
            bullet_text=request.form['bullet_text'],
            template_text=request.form.get('template_text'),
            experience_id=exp_id,
            tags=request.form.get('tags', ''),
            category=request.form.get('category', '')
        )
        flash('Bullet added successfully', 'success')
        return redirect(url_for('resume.view_resume'))
    
    experiences = Experience.get_all()
    return render_template('add_bullet.html', experiences=experiences)


@bp.route('/bullet/<int:bullet_id>/edit', methods=['GET', 'POST'])
def edit_bullet(bullet_id):
    """Edit a bullet"""
    if request.method == 'POST':
        if Bullet.update(
            bullet_id,
            bullet_text=request.form['bullet_text'],
            template_text=request.form.get('template_text', ''),
            tags=request.form.get('tags', ''),
            category=request.form.get('category', '')
        ):
            flash('Bullet updated successfully', 'success')
        else:
            flash('Bullet not found', 'error')
        return redirect(url_for('resume.view_resume'))
    
    bullet = Bullet.get_by_id(bullet_id)
    if not bullet:
        flash('Bullet not found', 'error')
        return redirect(url_for('resume.view_resume'))
    
    return render_template('edit_bullet.html', bullet=bullet)


@bp.route('/bullet/<int:bullet_id>/delete', methods=['POST'])
def delete_bullet(bullet_id):
    """Delete a bullet"""
    if Bullet.delete(bullet_id):
        flash('Bullet deleted successfully', 'success')
    else:
        flash('Bullet not found', 'error')
    return redirect(url_for('resume.view_resume') + '#bullets-section')


# ============================================================================
# SKILL CRUD
# ============================================================================

@bp.route('/skill/add', methods=['GET', 'POST'])
def add_skill():
    """Add a new skill"""
    if request.method == 'POST':
        Skill.create(
            skill_name=request.form['skill_name'],
            category=request.form.get('category', '')
        )
        flash('Skill added successfully', 'success')
        return redirect(url_for('resume.view_resume'))
    
    return render_template('add_skill.html')


@bp.route('/skill/<int:skill_id>/edit', methods=['GET', 'POST'])
def edit_skill(skill_id):
    """Edit a skill"""
    if request.method == 'POST':
        if Skill.update(
            skill_id,
            skill_name=request.form['skill_name'],
            category=request.form.get('category', '')
        ):
            flash('Skill updated successfully', 'success')
        else:
            flash('Skill not found', 'error')
        return redirect(url_for('resume.view_resume'))
    
    skill = Skill.get_by_id(skill_id)
    if not skill:
        flash('Skill not found', 'error')
        return redirect(url_for('resume.view_resume'))
    
    return render_template('edit_skill.html', skill=skill)


@bp.route('/skill/<int:skill_id>/delete', methods=['POST'])
def delete_skill(skill_id):
    """Delete a skill"""
    if Skill.delete(skill_id):
        flash('Skill deleted successfully', 'success')
    else:
        flash('Skill not found', 'error')
    return redirect(url_for('resume.view_resume') + '#skills-section')


# ============================================================================
# EDUCATION CRUD
# ============================================================================

@bp.route('/education/add', methods=['GET', 'POST'])
def add_education():
    """Add a new education entry"""
    if request.method == 'POST':
        Education.create(
            school_name=request.form['school'],
            degree=request.form.get('degree', ''),
            field_of_study=request.form.get('field', ''),
            graduation_year=request.form.get('grad_year', ''),
            location=request.form.get('location', '')
        )
        flash('Education added successfully', 'success')
        return redirect(url_for('resume.view_resume'))
    
    return render_template('add_education.html')


@bp.route('/education/<int:edu_id>/edit', methods=['GET', 'POST'])
def edit_education(edu_id):
    """Edit an education entry"""
    if request.method == 'POST':
        if Education.update(
            edu_id,
            school_name=request.form['school'],
            degree=request.form.get('degree', ''),
            field_of_study=request.form.get('field', ''),
            graduation_year=request.form.get('grad_year', ''),
            location=request.form.get('location', '')
        ):
            flash('Education updated successfully', 'success')
        else:
            flash('Education not found', 'error')
        return redirect(url_for('resume.view_resume'))
    
    education = Education.get_by_id(edu_id)
    if not education:
        flash('Education not found', 'error')
        return redirect(url_for('resume.view_resume'))
    
    return render_template('edit_education.html', edu=education)


@bp.route('/education/<int:edu_id>/delete', methods=['POST'])
def delete_education(edu_id):
    """Delete an education entry"""
    if Education.delete(edu_id):
        flash('Education deleted successfully', 'success')
    else:
        flash('Education not found', 'error')
    return redirect(url_for('resume.view_resume') + '#education-section')


# ============================================================================
# BULK DELETE OPERATIONS
# ============================================================================

@bp.route('/delete-section/<section_type>', methods=['POST'])
def delete_section(section_type):
    """Delete all components in a specific section"""
    section_map = {
        'experiences': (Experience, 'experiences-section'),
        'bullets': (Bullet, 'bullets-section'),
        'skills': (Skill, 'skills-section'),
        'education': (Education, 'education-section')
    }
    
    if section_type not in section_map:
        flash('Invalid section type', 'error')
        return redirect(url_for('resume.view_resume'))
    
    model_class, anchor = section_map[section_type]
    
    try:
        count = model_class.delete_all()
        flash(f'Deleted all {count} {section_type}', 'warning')
        return redirect(url_for('resume.view_resume') + f'#{anchor}')
    except Exception as e:
        flash(f'Error deleting {section_type}: {str(e)}', 'error')
        return redirect(url_for('resume.view_resume'))

@bp.route('/delete-all', methods=['POST'])
def delete_all_components():
    """Delete all resume components"""
    confirm_text = request.form.get('confirm_delete_all', '')
    
    if confirm_text != 'DELETE_EVERYTHING':
        flash('Confirmation text did not match. Nothing was deleted.', 'error')
        return redirect(url_for('resume.view_resume'))
    
    from models.database import get_db_context
    
    try:
        with get_db_context() as (conn, cursor):
            # Count items before deletion
            cursor.execute("SELECT COUNT(*) FROM experiences")
            exp_count = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM bullets")
            bullet_count = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM skills")
            skill_count = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM education")
            edu_count = cursor.fetchone()[0]
            
            total = exp_count + bullet_count + skill_count + edu_count
            
            # Delete everything
            cursor.execute("DELETE FROM experiences")
            cursor.execute("DELETE FROM bullets")
            cursor.execute("DELETE FROM skills")
            cursor.execute("DELETE FROM education")
            
            conn.commit()
        
        flash(
            f'Successfully deleted all resume components: {exp_count} experiences, '
            f'{bullet_count} bullets, {skill_count} skills, {edu_count} education entries',
            'success'
        )
    except Exception as e:
        flash(f'Error deleting components: {str(e)}', 'error')

    return redirect(url_for('resume.view_resume'))


@bp.route('/skills/cleanup-preview', methods=['GET'])
def cleanup_skills_preview():
    """Generate AI suggestions for cleaning up skill categories"""
    skills = Skill.get_all()

    if not skills:
        flash('No skills to clean up', 'info')
        return redirect(url_for('resume.view_resume'))

    try:
        ai = get_ai_service()
        cleanup_suggestions = ai.cleanup_skill_categories(skills)

        # Store in session for the apply step
        session['skill_cleanup_suggestions'] = json.dumps(cleanup_suggestions)

        return render_template(
            'skill_cleanup_preview.html',
            suggestions=cleanup_suggestions,
            skills=skills
        )
    except Exception as e:
        flash(f'Error generating cleanup suggestions: {str(e)}', 'error')
        return redirect(url_for('resume.view_resume'))


@bp.route('/skills/cleanup-apply', methods=['POST'])
def cleanup_skills_apply():
    """Apply the AI-suggested category cleanup"""
    if 'skill_cleanup_suggestions' not in session:
        flash('No cleanup suggestions found. Please try again.', 'error')
        return redirect(url_for('resume.view_resume'))

    try:
        suggestions = json.loads(session['skill_cleanup_suggestions'])
        category_mappings = suggestions.get('category_mappings', [])

        updated_count = 0
        with get_db_context() as (conn, cursor):
            for mapping in category_mappings:
                old_category = mapping['old_category']
                new_category = mapping['new_category']

                # Update all skills with the old category to the new category
                cursor.execute(
                    "UPDATE skills SET category = ? WHERE category = ?",
                    (new_category, old_category)
                )
                updated_count += cursor.rowcount

            conn.commit()

        # Clear session
        session.pop('skill_cleanup_suggestions', None)

        flash(
            f'Successfully cleaned up skill categories! Updated {updated_count} skill(s).',
            'success'
        )
    except Exception as e:
        flash(f'Error applying cleanup: {str(e)}', 'error')

    return redirect(url_for('resume.view_resume') + '#skills-section')


# ============================================================================
# EXPORT PROFILES
# ============================================================================

@bp.route('/export/profiles')
def list_profiles():
    """List all export profiles"""
    profiles = ExportProfile.get_all_with_rule_counts()
    return render_template('export_profiles.html', profiles=profiles)


@bp.route('/export/profiles/new', methods=['GET', 'POST'])
def create_profile():
    """Create a new export profile"""
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        description = request.form.get('description', '').strip()

        if not name:
            flash('Profile name is required', 'error')
            return render_template('create_export_profile.html')

        profile_id = ExportProfile.create(name, description)
        flash(f'Profile "{name}" created!', 'success')
        return redirect(url_for('resume.edit_profile', profile_id=profile_id))

    return render_template('create_export_profile.html')


@bp.route('/export/profiles/<int:profile_id>/edit', methods=['GET', 'POST'])
def edit_profile(profile_id):
    """Edit an export profile and manage its rules"""
    profile = ExportProfile.get_by_id(profile_id)
    if not profile:
        flash('Profile not found', 'error')
        return redirect(url_for('resume.list_profiles'))

    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        description = request.form.get('description', '').strip()

        if not name:
            flash('Profile name is required', 'error')
        else:
            ExportProfile.update(profile_id, name, description)
            flash('Profile updated!', 'success')
            return redirect(url_for('resume.edit_profile', profile_id=profile_id))

    rules = ExportProfile.get_rules(profile_id)

    # Parse rule configs and add descriptions for display
    parsed_rules = []
    for rule in rules:
        config = json.loads(rule['config'])
        parsed_rules.append({
            'id': rule['id'],
            'rule_type': rule['rule_type'],
            'rule_type_label': ExportProfile.RULE_TYPE_LABELS.get(rule['rule_type'], rule['rule_type']),
            'rule_order': rule['rule_order'],
            'config': config,
            'enabled': bool(rule['enabled']),
            'description': ExportProfile.describe_rule(rule['rule_type'], config),
        })

    # Gather data for rule config forms
    with get_db_context() as (conn, cursor):
        cursor.execute("SELECT DISTINCT category FROM skills WHERE category IS NOT NULL ORDER BY category")
        skill_categories = [row['category'] for row in cursor.fetchall()]

        cursor.execute("SELECT id, skill_name, category FROM skills ORDER BY category, skill_name")
        all_skills = cursor.fetchall()

        cursor.execute("SELECT DISTINCT category FROM bullets WHERE category IS NOT NULL ORDER BY category")
        bullet_categories = [row['category'] for row in cursor.fetchall()]

        cursor.execute("SELECT id, job_title, company_name, alternate_titles FROM experiences ORDER BY id")
        experiences = cursor.fetchall()

    # Build skills-by-category dict for the split rule form
    skills_by_category = {}
    for skill in all_skills:
        cat = skill['category'] or 'General'
        if cat not in skills_by_category:
            skills_by_category[cat] = []
        skills_by_category[cat].append({'id': skill['id'], 'name': skill['skill_name']})

    # Load header info
    header_info = ExportProfile.get_header_info(profile_id)

    return render_template(
        'edit_export_profile.html',
        profile=profile,
        rules=parsed_rules,
        header_info=header_info,
        skill_categories=skill_categories,
        bullet_categories=bullet_categories,
        experiences=experiences,
        skills_by_category_json=json.dumps(skills_by_category),
        rule_types=ExportProfile.RULE_TYPES,
        rule_type_labels=ExportProfile.RULE_TYPE_LABELS,
    )


@bp.route('/export/profiles/<int:profile_id>/header', methods=['POST'])
def update_header_info(profile_id):
    """Update the personal header info for a profile"""
    profile = ExportProfile.get_by_id(profile_id)
    if not profile:
        flash('Profile not found', 'error')
        return redirect(url_for('resume.list_profiles'))

    header_info = {
        'name': request.form.get('header_name', '').strip(),
        'email': request.form.get('header_email', '').strip(),
        'phone': request.form.get('header_phone', '').strip(),
        'location': request.form.get('header_location', '').strip(),
        'links': request.form.get('header_links', '').strip(),
    }

    # Remove empty fields so the JSON stays clean
    header_info = {k: v for k, v in header_info.items() if v}

    ExportProfile.update_header_info(profile_id, header_info)
    flash('Personal header updated!', 'success')
    return redirect(url_for('resume.edit_profile', profile_id=profile_id))


@bp.route('/export/profiles/<int:profile_id>/delete', methods=['POST'])
def delete_profile(profile_id):
    """Delete an export profile"""
    profile = ExportProfile.get_by_id(profile_id)
    if profile:
        ExportProfile.delete(profile_id)
        flash(f'Profile "{profile["name"]}" deleted', 'success')
    else:
        flash('Profile not found', 'error')
    return redirect(url_for('resume.list_profiles'))


@bp.route('/export/profiles/<int:profile_id>/duplicate', methods=['POST'])
def duplicate_profile(profile_id):
    """Duplicate an export profile"""
    profile = ExportProfile.get_by_id(profile_id)
    if not profile:
        flash('Profile not found', 'error')
        return redirect(url_for('resume.list_profiles'))

    new_name = f"{profile['name']} (Copy)"
    new_id = ExportProfile.duplicate(profile_id, new_name)
    flash(f'Profile duplicated as "{new_name}"', 'success')
    return redirect(url_for('resume.edit_profile', profile_id=new_id))


@bp.route('/export/profiles/<int:profile_id>/set-default', methods=['POST'])
def set_default_profile(profile_id):
    """Set a profile as the default"""
    profile = ExportProfile.get_by_id(profile_id)
    if not profile:
        flash('Profile not found', 'error')
        return redirect(url_for('resume.list_profiles'))

    ExportProfile.set_default(profile_id)
    flash(f'"{profile["name"]}" set as default profile', 'success')
    return redirect(url_for('resume.list_profiles'))


@bp.route('/export/profiles/<int:profile_id>/clear-default', methods=['POST'])
def clear_default_profile(profile_id):
    """Clear default status from a profile"""
    ExportProfile.clear_default()
    flash('Default profile cleared', 'success')
    return redirect(url_for('resume.list_profiles'))


@bp.route('/export/profiles/<int:profile_id>/rules/add', methods=['POST'])
def add_rule(profile_id):
    """Add a rule to a profile"""
    rule_type = request.form.get('rule_type', '')
    if rule_type not in ExportProfile.RULE_TYPES:
        flash('Invalid rule type', 'error')
        return redirect(url_for('resume.edit_profile', profile_id=profile_id))

    # Build config from form data based on rule type
    config = _build_rule_config(rule_type, request.form)
    if config is None:
        flash('Invalid rule configuration', 'error')
        return redirect(url_for('resume.edit_profile', profile_id=profile_id))

    ExportProfile.add_rule(profile_id, rule_type, config)
    flash('Rule added!', 'success')
    return redirect(url_for('resume.edit_profile', profile_id=profile_id))


@bp.route('/export/profiles/rules/<int:rule_id>/delete', methods=['POST'])
def delete_rule(rule_id):
    """Delete a rule"""
    rule = ExportProfile.get_rule_by_id(rule_id)
    if not rule:
        flash('Rule not found', 'error')
        return redirect(url_for('resume.list_profiles'))

    profile_id = rule['profile_id']
    ExportProfile.delete_rule(rule_id)
    flash('Rule deleted', 'success')
    return redirect(url_for('resume.edit_profile', profile_id=profile_id))


@bp.route('/export/profiles/rules/<int:rule_id>/toggle', methods=['POST'])
def toggle_rule(rule_id):
    """Toggle a rule's enabled state"""
    rule = ExportProfile.get_rule_by_id(rule_id)
    if not rule:
        flash('Rule not found', 'error')
        return redirect(url_for('resume.list_profiles'))

    profile_id = rule['profile_id']
    ExportProfile.toggle_rule(rule_id)
    return redirect(url_for('resume.edit_profile', profile_id=profile_id))


def _build_rule_config(rule_type, form):
    """Build a config dict from form data based on rule type"""
    if rule_type == ExportProfile.RULE_RENAME_CATEGORY:
        target = form.get('rename_target', 'skills')
        from_name = form.get('rename_from', '').strip()
        to_name = form.get('rename_to', '').strip()
        if not from_name or not to_name:
            return None
        return {'target': target, 'from_name': from_name, 'to_name': to_name}

    elif rule_type == ExportProfile.RULE_MERGE_CATEGORIES:
        target = form.get('merge_target', 'skills')
        source_categories = form.getlist('merge_sources')
        destination = form.get('merge_destination', '').strip()
        if not source_categories or not destination:
            return None
        return {'target': target, 'source_categories': source_categories,
                'destination_category': destination}

    elif rule_type == ExportProfile.RULE_SPLIT_CATEGORY:
        source_category = form.get('split_source', '').strip()
        # Parse splits: each split has a new_category name and a list of skill_ids
        splits = []
        split_count = int(form.get('split_count', 0))
        for i in range(split_count):
            new_cat = form.get(f'split_name_{i}', '').strip()
            skill_ids_str = form.getlist(f'split_skills_{i}')
            skill_ids = [int(sid) for sid in skill_ids_str if sid.isdigit()]
            if new_cat and skill_ids:
                splits.append({'new_category': new_cat, 'skill_ids': skill_ids})
        if not source_category or not splits:
            return None
        return {'target': 'skills', 'source_category': source_category, 'splits': splits}

    elif rule_type == ExportProfile.RULE_SECTION_ORDER:
        order = form.getlist('section_order')
        if not order:
            return None
        return {'order': order}

    elif rule_type == ExportProfile.RULE_USE_ALTERNATE_TITLE:
        exp_id_str = form.get('alt_title_exp_id', '')
        title = form.get('alt_title_value', '').strip()
        if not exp_id_str.isdigit() or not title:
            return None
        return {'experience_id': int(exp_id_str), 'title': title}

    return None


# ============================================================================
# EXPORT
# ============================================================================

@bp.route('/export')
def export_select():
    """Show component selection page for resume export"""
    with get_db_context() as (conn, cursor):
        # Get all resume components
        cursor.execute("SELECT * FROM experiences ORDER BY id")
        experiences = cursor.fetchall()

        cursor.execute("SELECT * FROM bullets ORDER BY experience_id, id")
        bullets = cursor.fetchall()

        cursor.execute("SELECT * FROM skills ORDER BY category, skill_name")
        skills = cursor.fetchall()

        cursor.execute("SELECT * FROM education ORDER BY id")
        education = cursor.fetchall()

    # Check if user has resume data
    if not experiences and not bullets and not skills and not education:
        flash('Please add resume components before exporting', 'error')
        return redirect(url_for('resume.view_resume'))

    # Group bullets by experience
    bullets_by_exp = {}
    for bullet in bullets:
        exp_id = bullet['experience_id']
        if exp_id not in bullets_by_exp:
            bullets_by_exp[exp_id] = []
        bullets_by_exp[exp_id].append(bullet)

    # Load export profiles
    profiles = ExportProfile.get_all_with_rule_counts()
    default_profile = ExportProfile.get_default()

    # Build profile rules data for JS (so rule toggles render client-side)
    profiles_rules_data = {}
    for profile in profiles:
        profile_data = ExportProfile.get_profile_with_rules(profile['id'])
        if profile_data:
            profiles_rules_data[profile['id']] = [
                {
                    'id': r['id'],
                    'description': ExportProfile.describe_rule(r['rule_type'], r['config']),
                    'enabled': r['enabled'],
                }
                for r in profile_data['rules']
            ]

    return render_template(
        'export_select.html',
        experiences=experiences,
        bullets_by_exp=bullets_by_exp,
        skills=skills,
        education=education,
        profiles=profiles,
        default_profile=default_profile,
        profiles_rules_json=json.dumps(profiles_rules_data),
    )


@bp.route('/export/generate', methods=['POST'])
def export_generate():
    """Generate and download resume based on selected components"""
    # Get selected component IDs from form
    selected_exp_ids = request.form.getlist('experience_ids')
    selected_bullet_ids = request.form.getlist('bullet_ids')
    selected_skill_ids = request.form.getlist('skill_ids')
    selected_edu_ids = request.form.getlist('education_ids')
    export_format = request.form.get('export_format', 'txt')

    with get_db_context() as (conn, cursor):
        # Fetch selected experiences
        experiences = []
        if selected_exp_ids:
            placeholders = ','.join('?' * len(selected_exp_ids))
            cursor.execute(f"SELECT * FROM experiences WHERE id IN ({placeholders}) ORDER BY id", selected_exp_ids)
            experiences = cursor.fetchall()

        # Fetch selected bullets
        bullets = []
        if selected_bullet_ids:
            placeholders = ','.join('?' * len(selected_bullet_ids))
            cursor.execute(f"SELECT * FROM bullets WHERE id IN ({placeholders}) ORDER BY experience_id, id", selected_bullet_ids)
            bullets = cursor.fetchall()

        # Fetch selected skills
        skills = []
        if selected_skill_ids:
            placeholders = ','.join('?' * len(selected_skill_ids))
            cursor.execute(f"SELECT * FROM skills WHERE id IN ({placeholders}) ORDER BY category, skill_name", selected_skill_ids)
            skills = cursor.fetchall()

        # Fetch selected education
        education = []
        if selected_edu_ids:
            placeholders = ','.join('?' * len(selected_edu_ids))
            cursor.execute(f"SELECT * FROM education WHERE id IN ({placeholders}) ORDER BY id", selected_edu_ids)
            education = cursor.fetchall()

    # Create filename
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

    # Apply export profile rules if a profile is selected
    profile_id = request.form.get('profile_id', '')
    disabled_rule_ids = request.form.getlist('disabled_rule_ids')
    disabled_rule_ids = {int(rid) for rid in disabled_rule_ids if rid.isdigit()}
    section_order = None
    header_info = None

    if profile_id and profile_id.isdigit():
        profile_data = ExportProfile.get_profile_with_rules(int(profile_id))
        if profile_data:
            # Filter out disabled rules (both profile-disabled and per-export-disabled)
            active_rules = [
                r for r in profile_data['rules']
                if r['enabled'] and r['id'] not in disabled_rule_ids
            ]

            # Convert Row objects to dicts for transformation
            exp_dicts = [dict(e) for e in experiences]
            bullet_dicts = [dict(b) for b in bullets]
            skill_dicts = [dict(s) for s in skills]
            edu_dicts = [dict(e) for e in education]

            transformed = apply_export_rules(exp_dicts, bullet_dicts, skill_dicts, edu_dicts, active_rules)
            experiences = transformed['experiences']
            bullets = transformed['bullets']
            skills = transformed['skills']
            education = transformed['education']
            section_order = transformed['section_order']

            # Extract personal header info from profile
            if profile_data.get('header_info'):
                header_info = profile_data['header_info']

    # Generate resume in requested format
    if export_format == 'txt':
        content_bytes, mimetype, filename = generate_resume_text(experiences, bullets, skills, education, timestamp, section_order, header_info)
    elif export_format == 'md':
        content_bytes, mimetype, filename = generate_resume_markdown(experiences, bullets, skills, education, timestamp, section_order, header_info)
    elif export_format == 'html':
        content_bytes, mimetype, filename = generate_resume_html(experiences, bullets, skills, education, timestamp, section_order, header_info)
    elif export_format == 'docx':
        content_bytes, mimetype, filename = generate_resume_docx(experiences, bullets, skills, education, timestamp, section_order, header_info)
    elif export_format == 'pdf':
        content_bytes, mimetype, filename = generate_resume_pdf(experiences, bullets, skills, education, timestamp, section_order, header_info)
    else:
        flash('Invalid export format', 'error')
        return redirect(url_for('resume.export_select'))

    # Return as download
    return send_file(
        content_bytes,
        mimetype=mimetype,
        as_attachment=True,
        download_name=filename
    )


def generate_resume_text(experiences, bullets, skills, education, timestamp, section_order=None, header_info=None):
    """Generate plain text resume from selected components"""
    if section_order is None:
        section_order = ['experience', 'skills', 'education']

    output = []
    output.append("=" * 60)

    if header_info and header_info.get('name'):
        output.append(header_info['name'].upper())
        contact_parts = []
        if header_info.get('email'):
            contact_parts.append(header_info['email'])
        if header_info.get('phone'):
            contact_parts.append(header_info['phone'])
        if header_info.get('location'):
            contact_parts.append(header_info['location'])
        if contact_parts:
            output.append(' | '.join(contact_parts))
        if header_info.get('links'):
            output.append(header_info['links'])
    else:
        output.append("RESUME")

    output.append("=" * 60)
    output.append("")

    def _render_experience():
        if not experiences:
            return
        output.append("WORK EXPERIENCE")
        output.append("-" * 60)
        output.append("")

        bullets_by_exp = {}
        for bullet in bullets:
            exp_id = bullet['experience_id']
            if exp_id not in bullets_by_exp:
                bullets_by_exp[exp_id] = []
            bullets_by_exp[exp_id].append(bullet)

        for exp in experiences:
            output.append(f"{exp['job_title']}")
            output.append(f"{exp['company_name']}")
            output.append(f"{exp['start_date']} - {exp['end_date']}")
            if exp['location']:
                output.append(f"{exp['location']}")
            output.append("")

            if exp['description']:
                output.append(exp['description'])
                output.append("")

            if exp['id'] in bullets_by_exp:
                for bullet in bullets_by_exp[exp['id']]:
                    output.append(f"  • {bullet['bullet_text']}")
                output.append("")

        output.append("")

    def _render_skills():
        if not skills:
            return
        output.append("SKILLS")
        output.append("-" * 60)
        output.append("")

        skills_by_cat = {}
        for skill in skills:
            cat = skill['category'] or 'General'
            if cat not in skills_by_cat:
                skills_by_cat[cat] = []
            skills_by_cat[cat].append(skill['skill_name'])

        for category, skill_list in skills_by_cat.items():
            output.append(f"{category}: {', '.join(skill_list)}")

        output.append("")
        output.append("")

    def _render_education():
        if not education:
            return
        output.append("EDUCATION")
        output.append("-" * 60)
        output.append("")

        for edu in education:
            output.append(f"{edu['degree']} in {edu['field_of_study']}")
            output.append(f"{edu['school_name']}")
            output.append(f"{edu['graduation_year']}")
            if edu['location']:
                output.append(f"{edu['location']}")
            output.append("")

    section_renderers = {
        'experience': _render_experience,
        'skills': _render_skills,
        'education': _render_education,
    }

    for section in section_order:
        renderer = section_renderers.get(section)
        if renderer:
            renderer()

    content = '\n'.join(output)
    content_bytes = BytesIO(content.encode('utf-8'))
    content_bytes.seek(0)

    return content_bytes, 'text/plain', f"resume_{timestamp}.txt"


def generate_resume_markdown(experiences, bullets, skills, education, timestamp, section_order=None, header_info=None):
    """Generate Markdown resume from selected components"""
    if section_order is None:
        section_order = ['experience', 'skills', 'education']

    output = []

    if header_info and header_info.get('name'):
        output.append(f"# {header_info['name']}")
        contact_parts = []
        if header_info.get('email'):
            contact_parts.append(header_info['email'])
        if header_info.get('phone'):
            contact_parts.append(header_info['phone'])
        if header_info.get('location'):
            contact_parts.append(header_info['location'])
        if contact_parts:
            output.append(' | '.join(contact_parts))
        if header_info.get('links'):
            output.append(header_info['links'])
    else:
        output.append("# RESUME")

    output.append("")

    def _render_experience():
        if not experiences:
            return
        output.append("## WORK EXPERIENCE")
        output.append("")

        bullets_by_exp = {}
        for bullet in bullets:
            exp_id = bullet['experience_id']
            if exp_id not in bullets_by_exp:
                bullets_by_exp[exp_id] = []
            bullets_by_exp[exp_id].append(bullet)

        for exp in experiences:
            output.append(f"### {exp['job_title']}")
            output.append(f"**{exp['company_name']}**")
            output.append(f"*{exp['start_date']} - {exp['end_date']}*")
            if exp['location']:
                output.append(f"*{exp['location']}*")
            output.append("")

            if exp['description']:
                output.append(exp['description'])
                output.append("")

            if exp['id'] in bullets_by_exp:
                for bullet in bullets_by_exp[exp['id']]:
                    output.append(f"- {bullet['bullet_text']}")
                output.append("")

    def _render_skills():
        if not skills:
            return
        output.append("## SKILLS")
        output.append("")

        skills_by_cat = {}
        for skill in skills:
            cat = skill['category'] or 'General'
            if cat not in skills_by_cat:
                skills_by_cat[cat] = []
            skills_by_cat[cat].append(skill['skill_name'])

        for category, skill_list in skills_by_cat.items():
            output.append(f"**{category}:** {', '.join(skill_list)}")
        output.append("")

    def _render_education():
        if not education:
            return
        output.append("## EDUCATION")
        output.append("")

        for edu in education:
            output.append(f"### {edu['degree']} in {edu['field_of_study']}")
            output.append(f"**{edu['school_name']}**")
            output.append(f"*{edu['graduation_year']}*")
            if edu['location']:
                output.append(f"*{edu['location']}*")
            output.append("")

    section_renderers = {
        'experience': _render_experience,
        'skills': _render_skills,
        'education': _render_education,
    }

    for section in section_order:
        renderer = section_renderers.get(section)
        if renderer:
            renderer()

    content = '\n'.join(output)
    content_bytes = BytesIO(content.encode('utf-8'))
    content_bytes.seek(0)

    return content_bytes, 'text/markdown', f"resume_{timestamp}.md"


def generate_resume_html(experiences, bullets, skills, education, timestamp, section_order=None, header_info=None):
    """Generate HTML resume from selected components"""
    if section_order is None:
        section_order = ['experience', 'skills', 'education']

    output = []
    output.append("<!DOCTYPE html>")
    output.append("<html lang='en'>")
    output.append("<head>")
    output.append("    <meta charset='UTF-8'>")
    output.append("    <meta name='viewport' content='width=device-width, initial-scale=1.0'>")
    output.append("    <title>Resume</title>")
    output.append("    <style>")
    output.append("        body { font-family: Arial, sans-serif; max-width: 800px; margin: 40px auto; padding: 20px; line-height: 1.6; }")
    output.append("        h1 { border-bottom: 3px solid #333; padding-bottom: 10px; }")
    output.append("        h2 { color: #333; border-bottom: 2px solid #666; padding-bottom: 5px; margin-top: 30px; }")
    output.append("        h3 { color: #555; margin-bottom: 5px; }")
    output.append("        .meta { color: #666; font-style: italic; margin: 5px 0; }")
    output.append("        ul { margin: 10px 0; }")
    output.append("        .header-contact { text-align: center; color: #555; margin: 5px 0; }")
    output.append("        .footer { text-align: center; margin-top: 40px; padding-top: 20px; border-top: 1px solid #ccc; color: #666; font-size: 0.9em; }")
    output.append("    </style>")
    output.append("</head>")
    output.append("<body>")

    if header_info and header_info.get('name'):
        output.append(f"    <h1 style='text-align: center;'>{html_escape(header_info['name'])}</h1>")
        contact_parts = []
        if header_info.get('email'):
            contact_parts.append(html_escape(header_info['email']))
        if header_info.get('phone'):
            contact_parts.append(html_escape(header_info['phone']))
        if header_info.get('location'):
            contact_parts.append(html_escape(header_info['location']))
        if contact_parts:
            output.append(f"    <p class='header-contact'>{' | '.join(contact_parts)}</p>")
        if header_info.get('links'):
            output.append(f"    <p class='header-contact'>{html_escape(header_info['links'])}</p>")
    else:
        output.append("    <h1>RESUME</h1>")

    def _render_experience():
        if not experiences:
            return
        output.append("    <h2>WORK EXPERIENCE</h2>")

        bullets_by_exp = {}
        for bullet in bullets:
            exp_id = bullet['experience_id']
            if exp_id not in bullets_by_exp:
                bullets_by_exp[exp_id] = []
            bullets_by_exp[exp_id].append(bullet)

        for exp in experiences:
            output.append(f"    <h3>{exp['job_title']}</h3>")
            output.append(f"    <div class='meta'><strong>{exp['company_name']}</strong></div>")
            output.append(f"    <div class='meta'>{exp['start_date']} - {exp['end_date']}")
            if exp['location']:
                output.append(f" | {exp['location']}")
            output.append("</div>")

            if exp['description']:
                output.append(f"    <p>{exp['description']}</p>")

            if exp['id'] in bullets_by_exp:
                output.append("    <ul>")
                for bullet in bullets_by_exp[exp['id']]:
                    output.append(f"        <li>{bullet['bullet_text']}</li>")
                output.append("    </ul>")

    def _render_skills():
        if not skills:
            return
        output.append("    <h2>SKILLS</h2>")

        skills_by_cat = {}
        for skill in skills:
            cat = skill['category'] or 'General'
            if cat not in skills_by_cat:
                skills_by_cat[cat] = []
            skills_by_cat[cat].append(skill['skill_name'])

        for category, skill_list in skills_by_cat.items():
            output.append(f"    <p><strong>{category}:</strong> {', '.join(skill_list)}</p>")

    def _render_education():
        if not education:
            return
        output.append("    <h2>EDUCATION</h2>")

        for edu in education:
            output.append(f"    <h3>{edu['degree']} in {edu['field_of_study']}</h3>")
            output.append(f"    <div class='meta'><strong>{edu['school_name']}</strong></div>")
            output.append(f"    <div class='meta'>{edu['graduation_year']}")
            if edu['location']:
                output.append(f" | {edu['location']}")
            output.append("</div>")

    section_renderers = {
        'experience': _render_experience,
        'skills': _render_skills,
        'education': _render_education,
    }

    for section in section_order:
        renderer = section_renderers.get(section)
        if renderer:
            renderer()

    output.append("</body>")
    output.append("</html>")

    content = '\n'.join(output)
    content_bytes = BytesIO(content.encode('utf-8'))
    content_bytes.seek(0)

    return content_bytes, 'text/html', f"resume_{timestamp}.html"


def generate_resume_docx(experiences, bullets, skills, education, timestamp, section_order=None, header_info=None):
    """Generate DOCX resume from selected components"""
    if section_order is None:
        section_order = ['experience', 'skills', 'education']

    doc = Document()

    if header_info and header_info.get('name'):
        title = doc.add_heading(header_info['name'], 0)
        title.alignment = 1  # Center

        contact_parts = []
        if header_info.get('email'):
            contact_parts.append(header_info['email'])
        if header_info.get('phone'):
            contact_parts.append(header_info['phone'])
        if header_info.get('location'):
            contact_parts.append(header_info['location'])
        if contact_parts:
            p = doc.add_paragraph()
            p.alignment = 1
            p.add_run(' | '.join(contact_parts))

        if header_info.get('links'):
            p = doc.add_paragraph()
            p.alignment = 1
            p.add_run(header_info['links'])
    else:
        title = doc.add_heading('RESUME', 0)
        title.alignment = 1  # Center alignment

    def _render_experience():
        if not experiences:
            return
        doc.add_heading('WORK EXPERIENCE', level=1)

        bullets_by_exp = {}
        for bullet in bullets:
            exp_id = bullet['experience_id']
            if exp_id not in bullets_by_exp:
                bullets_by_exp[exp_id] = []
            bullets_by_exp[exp_id].append(bullet)

        for exp in experiences:
            doc.add_heading(exp['job_title'], level=2)

            p = doc.add_paragraph()
            run = p.add_run(exp['company_name'])
            run.bold = True

            p = doc.add_paragraph()
            run = p.add_run(f"{exp['start_date']} - {exp['end_date']}")
            run.italic = True
            if exp['location']:
                run2 = p.add_run(f" | {exp['location']}")
                run2.italic = True

            if exp['description']:
                doc.add_paragraph(exp['description'])

            if exp['id'] in bullets_by_exp:
                for bullet in bullets_by_exp[exp['id']]:
                    doc.add_paragraph(bullet['bullet_text'], style='List Bullet')

    def _render_skills():
        if not skills:
            return
        doc.add_heading('SKILLS', level=1)

        skills_by_cat = {}
        for skill in skills:
            cat = skill['category'] or 'General'
            if cat not in skills_by_cat:
                skills_by_cat[cat] = []
            skills_by_cat[cat].append(skill['skill_name'])

        for category, skill_list in skills_by_cat.items():
            p = doc.add_paragraph()
            run = p.add_run(f"{category}: ")
            run.bold = True
            p.add_run(', '.join(skill_list))

    def _render_education():
        if not education:
            return
        doc.add_heading('EDUCATION', level=1)

        for edu in education:
            doc.add_heading(f"{edu['degree']} in {edu['field_of_study']}", level=2)

            p = doc.add_paragraph()
            run = p.add_run(edu['school_name'])
            run.bold = True

            p = doc.add_paragraph()
            run = p.add_run(edu['graduation_year'])
            run.italic = True
            if edu['location']:
                run2 = p.add_run(f" | {edu['location']}")
                run2.italic = True

    section_renderers = {
        'experience': _render_experience,
        'skills': _render_skills,
        'education': _render_education,
    }

    for section in section_order:
        renderer = section_renderers.get(section)
        if renderer:
            renderer()


    # Save to BytesIO
    content_bytes = BytesIO()
    doc.save(content_bytes)
    content_bytes.seek(0)

    return content_bytes, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', f"resume_{timestamp}.docx"


def generate_resume_pdf(experiences, bullets, skills, education, timestamp, section_order=None, header_info=None):
    """Generate PDF resume from selected components"""
    if section_order is None:
        section_order = ['experience', 'skills', 'education']

    content_bytes = BytesIO()
    doc = SimpleDocTemplate(content_bytes, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        alignment=1,  # Center
        spaceAfter=6
    )

    contact_style = ParagraphStyle(
        'ContactInfo',
        parent=styles['Normal'],
        fontSize=10,
        alignment=1,  # Center
        spaceAfter=4
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=12
    )

    subheading_style = ParagraphStyle(
        'CustomSubHeading',
        parent=styles['Heading3'],
        fontSize=12,
        spaceAfter=6
    )

    # Title / Header
    if header_info and header_info.get('name'):
        story.append(Paragraph(header_info['name'], title_style))
        contact_parts = []
        if header_info.get('email'):
            contact_parts.append(header_info['email'])
        if header_info.get('phone'):
            contact_parts.append(header_info['phone'])
        if header_info.get('location'):
            contact_parts.append(header_info['location'])
        if contact_parts:
            story.append(Paragraph(' | '.join(contact_parts), contact_style))
        if header_info.get('links'):
            story.append(Paragraph(header_info['links'], contact_style))
    else:
        story.append(Paragraph("RESUME", title_style))

    story.append(Spacer(1, 0.2*inch))

    def _render_experience():
        if not experiences:
            return
        story.append(Paragraph("WORK EXPERIENCE", heading_style))

        bullets_by_exp = {}
        for bullet in bullets:
            exp_id = bullet['experience_id']
            if exp_id not in bullets_by_exp:
                bullets_by_exp[exp_id] = []
            bullets_by_exp[exp_id].append(bullet)

        for exp in experiences:
            story.append(Paragraph(exp['job_title'], subheading_style))
            story.append(Paragraph(f"<b>{exp['company_name']}</b>", styles['Normal']))

            location_date = f"<i>{exp['start_date']} - {exp['end_date']}"
            if exp['location']:
                location_date += f" | {exp['location']}"
            location_date += "</i>"
            story.append(Paragraph(location_date, styles['Normal']))

            if exp['description']:
                story.append(Spacer(1, 0.1*inch))
                story.append(Paragraph(exp['description'], styles['Normal']))

            if exp['id'] in bullets_by_exp:
                story.append(Spacer(1, 0.1*inch))
                for bullet in bullets_by_exp[exp['id']]:
                    story.append(Paragraph(f"&bull; {bullet['bullet_text']}", styles['Normal']))

            story.append(Spacer(1, 0.2*inch))

    def _render_skills():
        if not skills:
            return
        story.append(Paragraph("SKILLS", heading_style))

        skills_by_cat = {}
        for skill in skills:
            cat = skill['category'] or 'General'
            if cat not in skills_by_cat:
                skills_by_cat[cat] = []
            skills_by_cat[cat].append(skill['skill_name'])

        for category, skill_list in skills_by_cat.items():
            story.append(Paragraph(f"<b>{category}:</b> {', '.join(skill_list)}", styles['Normal']))

        story.append(Spacer(1, 0.2*inch))

    def _render_education():
        if not education:
            return
        story.append(Paragraph("EDUCATION", heading_style))

        for edu in education:
            story.append(Paragraph(f"{edu['degree']} in {edu['field_of_study']}", subheading_style))
            story.append(Paragraph(f"<b>{edu['school_name']}</b>", styles['Normal']))

            grad_location = f"<i>{edu['graduation_year']}"
            if edu['location']:
                grad_location += f" | {edu['location']}"
            grad_location += "</i>"
            story.append(Paragraph(grad_location, styles['Normal']))
            story.append(Spacer(1, 0.1*inch))

    section_renderers = {
        'experience': _render_experience,
        'skills': _render_skills,
        'education': _render_education,
    }

    for section in section_order:
        renderer = section_renderers.get(section)
        if renderer:
            renderer()

    # Build PDF
    doc.build(story)
    content_bytes.seek(0)

    return content_bytes, 'application/pdf', f"resume_{timestamp}.pdf"