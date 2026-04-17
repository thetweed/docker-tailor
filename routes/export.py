"""
Export Routes - Resume export, export profiles, and format generation
"""
import json
from datetime import datetime
from html import escape as html_escape
from io import BytesIO
from flask import Blueprint, render_template, request, redirect, url_for, flash, send_file
from models import Experience, Bullet, Skill, Education, ExportProfile
from models.database import get_db_context
from models.resume import get_all_components
from utils.json_helpers import safe_json_loads
from services.export_transform import apply_export_rules
from docx import Document
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

bp = Blueprint('export', __name__, url_prefix='/export')


def _require_profile(profile_id):
    """Return the profile row, or (None, redirect_response) if not found."""
    profile = ExportProfile.get_by_id(profile_id)
    if not profile:
        flash('Profile not found', 'error')
        return None, redirect(url_for('export.export_home'))
    return profile, None


# ============================================================================
# EXPORT HOME
# ============================================================================

@bp.route('/')
def export_home():
    """Export landing page with profiles and quick actions"""
    profiles = ExportProfile.get_all_with_rule_counts()
    return render_template('export.html', profiles=profiles)


# ============================================================================
# EXPORT PROFILES
# ============================================================================

@bp.route('/profiles/new', methods=['GET', 'POST'])
def create_profile():
    """Create a new export profile"""
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        description = request.form.get('description', '').strip()

        if not name:
            flash('Profile name is required', 'error')
            return render_template('create_export_profile.html')

        profile_id = ExportProfile.create(name, description)
        flash(f'Profile "{name}" created!', 'success')
        return redirect(url_for('export.edit_profile', profile_id=profile_id))

    return render_template('create_export_profile.html')


@bp.route('/profiles/<int:profile_id>/edit', methods=['GET', 'POST'])
def edit_profile(profile_id):
    """Edit an export profile and manage its rules"""
    profile, err = _require_profile(profile_id)
    if err:
        return err

    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        description = request.form.get('description', '').strip()

        if not name:
            flash('Profile name is required', 'error')
        else:
            ExportProfile.update(profile_id, name, description)
            flash('Profile updated!', 'success')
            return redirect(url_for('export.edit_profile', profile_id=profile_id))

    rules = ExportProfile.get_rules(profile_id)

    # Parse rule configs and add descriptions for display
    parsed_rules = []
    for rule in rules:
        config = safe_json_loads(rule['config'], f'rule id={rule["id"]}')
        parsed_rules.append({
            'id': rule['id'],
            'rule_type': rule['rule_type'],
            'rule_type_label': ExportProfile.RULE_TYPE_LABELS.get(rule['rule_type'], rule['rule_type']),
            'rule_order': rule['rule_order'],
            'config': config,
            'enabled': bool(rule['enabled']),
            'description': ExportProfile.describe_rule(rule['rule_type'], config),
        })

    # Gather data for rule config forms
    with get_db_context() as (conn, cursor):
        cursor.execute("SELECT id, skill_name, category FROM skills ORDER BY category, skill_name")
        all_skills = cursor.fetchall()

        cursor.execute("SELECT DISTINCT category FROM bullets WHERE category IS NOT NULL ORDER BY category")
        bullet_categories = [row['category'] for row in cursor.fetchall()]

        cursor.execute("SELECT id, job_title, company_name, alternate_titles FROM experiences ORDER BY id")
        experiences = cursor.fetchall()

    # Derive skill categories from the already-fetched skills list (preserves sorted order)
    seen_cats = set()
    skill_categories = []
    for skill in all_skills:
        cat = skill['category']
        if cat and cat not in seen_cats:
            seen_cats.add(cat)
            skill_categories.append(cat)

    # Build skills-by-category dict for the split rule form
    skills_by_category = {}
    for skill in all_skills:
        cat = skill['category'] or 'General'
        if cat not in skills_by_category:
            skills_by_category[cat] = []
        skills_by_category[cat].append({'id': skill['id'], 'name': skill['skill_name']})

    header_info = ExportProfile.parse_header_info(profile)

    return render_template(
        'edit_export_profile.html',
        profile=profile,
        rules=parsed_rules,
        header_info=header_info,
        skill_categories=skill_categories,
        bullet_categories=bullet_categories,
        experiences=experiences,
        skills_by_category_json=json.dumps(skills_by_category),
        rule_types=ExportProfile.RULE_TYPES,
        rule_type_labels=ExportProfile.RULE_TYPE_LABELS,
    )


@bp.route('/profiles/<int:profile_id>/header', methods=['POST'])
def update_header_info(profile_id):
    """Update the personal header info for a profile"""
    profile, err = _require_profile(profile_id)
    if err:
        return err

    header_info = {
        'name': request.form.get('header_name', '').strip(),
        'email': request.form.get('header_email', '').strip(),
        'phone': request.form.get('header_phone', '').strip(),
        'location': request.form.get('header_location', '').strip(),
        'links': request.form.get('header_links', '').strip(),
    }

    # Remove empty fields so the JSON stays clean
    header_info = {k: v for k, v in header_info.items() if v}

    ExportProfile.update_header_info(profile_id, header_info)
    flash('Personal header updated!', 'success')
    return redirect(url_for('export.edit_profile', profile_id=profile_id))


@bp.route('/profiles/<int:profile_id>/delete', methods=['POST'])
def delete_profile(profile_id):
    """Delete an export profile"""
    profile, err = _require_profile(profile_id)
    if err:
        return err
    ExportProfile.delete(profile_id)
    flash(f'Profile "{profile["name"]}" deleted', 'success')
    return redirect(url_for('export.export_home'))


@bp.route('/profiles/<int:profile_id>/duplicate', methods=['POST'])
def duplicate_profile(profile_id):
    """Duplicate an export profile"""
    profile, err = _require_profile(profile_id)
    if err:
        return err

    new_name = f"{profile['name']} (Copy)"
    new_id = ExportProfile.duplicate(profile_id, new_name)
    flash(f'Profile duplicated as "{new_name}"', 'success')
    return redirect(url_for('export.edit_profile', profile_id=new_id))


@bp.route('/profiles/<int:profile_id>/set-default', methods=['POST'])
def set_default_profile(profile_id):
    """Set a profile as the default"""
    profile, err = _require_profile(profile_id)
    if err:
        return err

    ExportProfile.set_default(profile_id)
    flash(f'"{profile["name"]}" set as default profile', 'success')
    return redirect(url_for('export.export_home'))


@bp.route('/profiles/<int:profile_id>/clear-default', methods=['POST'])
def clear_default_profile(profile_id):
    """Clear default status from a profile"""
    ExportProfile.clear_default()
    flash('Default profile cleared', 'success')
    return redirect(url_for('export.export_home'))


@bp.route('/profiles/<int:profile_id>/rules/add', methods=['POST'])
def add_rule(profile_id):
    """Add a rule to a profile"""
    rule_type = request.form.get('rule_type', '')
    if rule_type not in ExportProfile.RULE_TYPES:
        flash('Invalid rule type', 'error')
        return redirect(url_for('export.edit_profile', profile_id=profile_id))

    # Build config from form data based on rule type
    config = _build_rule_config(rule_type, request.form)
    if config is None:
        flash('Invalid rule configuration', 'error')
        return redirect(url_for('export.edit_profile', profile_id=profile_id))

    ExportProfile.add_rule(profile_id, rule_type, config)
    flash('Rule added!', 'success')
    return redirect(url_for('export.edit_profile', profile_id=profile_id))


@bp.route('/profiles/rules/<int:rule_id>/delete', methods=['POST'])
def delete_rule(rule_id):
    """Delete a rule"""
    rule = ExportProfile.get_rule_by_id(rule_id)
    if not rule:
        flash('Rule not found', 'error')
        return redirect(url_for('export.export_home'))

    profile_id = rule['profile_id']
    ExportProfile.delete_rule(rule_id)
    flash('Rule deleted', 'success')
    return redirect(url_for('export.edit_profile', profile_id=profile_id))


@bp.route('/profiles/rules/<int:rule_id>/toggle', methods=['POST'])
def toggle_rule(rule_id):
    """Toggle a rule's enabled state"""
    rule = ExportProfile.get_rule_by_id(rule_id)
    if not rule:
        flash('Rule not found', 'error')
        return redirect(url_for('export.export_home'))

    profile_id = rule['profile_id']
    ExportProfile.toggle_rule(rule_id)
    return redirect(url_for('export.edit_profile', profile_id=profile_id))


def _build_rule_config(rule_type, form):
    """Build a config dict from form data based on rule type"""
    if rule_type == ExportProfile.RULE_RENAME_CATEGORY:
        target = form.get('rename_target', 'skills')
        from_name = form.get('rename_from', '').strip()
        to_name = form.get('rename_to', '').strip()
        if not from_name or not to_name:
            return None
        return {'target': target, 'from_name': from_name, 'to_name': to_name}

    elif rule_type == ExportProfile.RULE_MERGE_CATEGORIES:
        target = form.get('merge_target', 'skills')
        source_categories = form.getlist('merge_sources')
        destination = form.get('merge_destination', '').strip()
        if not source_categories or not destination:
            return None
        return {'target': target, 'source_categories': source_categories,
                'destination_category': destination}

    elif rule_type == ExportProfile.RULE_SPLIT_CATEGORY:
        source_category = form.get('split_source', '').strip()
        # Parse splits: each split has a new_category name and a list of skill_ids
        splits = []
        try:
            split_count = int(form.get('split_count', 0))
        except (ValueError, TypeError):
            return None
        for i in range(split_count):
            new_cat = form.get(f'split_name_{i}', '').strip()
            skill_ids_str = form.getlist(f'split_skills_{i}')
            skill_ids = [int(sid) for sid in skill_ids_str if sid.isdigit()]
            if new_cat and skill_ids:
                splits.append({'new_category': new_cat, 'skill_ids': skill_ids})
        if not source_category or not splits:
            return None
        return {'target': 'skills', 'source_category': source_category, 'splits': splits}

    elif rule_type == ExportProfile.RULE_SECTION_ORDER:
        order = form.getlist('section_order')
        if not order:
            return None
        return {'order': order}

    elif rule_type == ExportProfile.RULE_USE_ALTERNATE_TITLE:
        exp_id_str = form.get('alt_title_exp_id', '')
        title = form.get('alt_title_value', '').strip()
        if not exp_id_str.isdigit() or not title:
            return None
        return {'experience_id': int(exp_id_str), 'title': title}

    elif rule_type == ExportProfile.RULE_RENAME_COMPANY:
        exp_id_str = form.get('rename_company_exp_id', '')
        display_name = form.get('rename_company_display_name', '').strip()
        if not exp_id_str.isdigit() or not display_name:
            return None
        return {'experience_id': int(exp_id_str), 'display_name': display_name}

    return None


# ============================================================================
# EXPORT SELECTION & GENERATION
# ============================================================================

@bp.route('/select')
def export_select():
    """Show component selection page for resume export"""
    experiences, bullets, skills, education = get_all_components()

    # Check if user has resume data
    if not experiences and not bullets and not skills and not education:
        flash('Please add resume components before exporting', 'error')
        return redirect(url_for('resume.view_resume'))

    # Group bullets by experience
    bullets_by_exp = {}
    for bullet in bullets:
        exp_id = bullet['experience_id']
        if exp_id not in bullets_by_exp:
            bullets_by_exp[exp_id] = []
        bullets_by_exp[exp_id].append(bullet)

    # Load export profiles
    profiles = ExportProfile.get_all_with_rule_counts()
    default_profile = ExportProfile.get_default()

    # Build profile rules data for JS (so rule toggles render client-side) — 2 queries, not N+1
    all_rules_by_profile = ExportProfile.get_all_rules_grouped()
    profiles_rules_data = {
        profile['id']: [
            {
                'id': r['id'],
                'description': ExportProfile.describe_rule(r['rule_type'], r['config']),
                'enabled': r['enabled'],
            }
            for r in all_rules_by_profile.get(profile['id'], [])
        ]
        for profile in profiles
    }

    # Check if pre-selecting from a tailor analysis
    analysis_id = request.args.get('analysis_id', type=int)
    pre_selected = None
    analysis_job = None

    if analysis_id:
        from models.tailor_analysis import TailorAnalysis
        recommended = TailorAnalysis.get_recommended_ids(analysis_id)

        if recommended:
            # Resolve any skill names that only have names (no valid IDs)
            skill_name_to_id = {s['skill_name'].lower(): s['id'] for s in skills}
            for name in recommended.get('skill_names', []):
                matched_id = skill_name_to_id.get(name.lower())
                if matched_id:
                    recommended['skill_ids'].add(matched_id)

            pre_selected = {
                'experience_ids': recommended['experience_ids'],
                'bullet_ids': recommended['bullet_ids'],
                'skill_ids': recommended['skill_ids'],
                'education_ids': recommended['education_ids'],
            }

            # Get job info for the context banner
            analysis = TailorAnalysis.get_by_id(analysis_id)
            if analysis:
                with get_db_context() as (conn, cursor):
                    cursor.execute("SELECT company_name, job_title FROM jobs WHERE id = ?",
                                   (analysis['job_id'],))
                    analysis_job = cursor.fetchone()

    return render_template(
        'export_select.html',
        experiences=experiences,
        bullets_by_exp=bullets_by_exp,
        skills=skills,
        education=education,
        profiles=profiles,
        default_profile=default_profile,
        profiles_rules_json=json.dumps(profiles_rules_data),
        pre_selected=pre_selected,
        analysis_job=analysis_job,
    )


@bp.route('/generate', methods=['POST'])
def export_generate():
    """Generate and download resume based on selected components"""
    # Get selected component IDs from form
    selected_exp_ids = request.form.getlist('experience_ids')
    selected_bullet_ids = request.form.getlist('bullet_ids')
    selected_skill_ids = request.form.getlist('skill_ids')
    selected_edu_ids = request.form.getlist('education_ids')
    export_format = request.form.get('export_format', 'txt')

    with get_db_context() as (conn, cursor):
        # Fetch selected experiences
        experiences = []
        if selected_exp_ids:
            placeholders = ','.join('?' * len(selected_exp_ids))
            cursor.execute(f"SELECT * FROM experiences WHERE id IN ({placeholders}) ORDER BY id", selected_exp_ids)
            experiences = cursor.fetchall()

        # Fetch selected bullets
        bullets = []
        if selected_bullet_ids:
            placeholders = ','.join('?' * len(selected_bullet_ids))
            cursor.execute(f"SELECT * FROM bullets WHERE id IN ({placeholders}) ORDER BY experience_id, id", selected_bullet_ids)
            bullets = cursor.fetchall()

        # Fetch selected skills
        skills = []
        if selected_skill_ids:
            placeholders = ','.join('?' * len(selected_skill_ids))
            cursor.execute(f"SELECT * FROM skills WHERE id IN ({placeholders}) ORDER BY category, skill_name", selected_skill_ids)
            skills = cursor.fetchall()

        # Fetch selected education
        education = []
        if selected_edu_ids:
            placeholders = ','.join('?' * len(selected_edu_ids))
            cursor.execute(f"SELECT * FROM education WHERE id IN ({placeholders}) ORDER BY id", selected_edu_ids)
            education = cursor.fetchall()

    # Create filename
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')

    # Apply export profile rules if a profile is selected
    profile_id = request.form.get('profile_id', '')
    disabled_rule_ids = request.form.getlist('disabled_rule_ids')
    disabled_rule_ids = {int(rid) for rid in disabled_rule_ids if rid.isdigit()}
    section_order = None
    header_info = None

    if profile_id and profile_id.isdigit():
        profile_data = ExportProfile.get_profile_with_rules(int(profile_id))
        if profile_data:
            # Filter out disabled rules (both profile-disabled and per-export-disabled)
            active_rules = [
                r for r in profile_data['rules']
                if r['enabled'] and r['id'] not in disabled_rule_ids
            ]

            # Convert Row objects to dicts for transformation
            exp_dicts = [dict(e) for e in experiences]
            bullet_dicts = [dict(b) for b in bullets]
            skill_dicts = [dict(s) for s in skills]
            edu_dicts = [dict(e) for e in education]

            transformed = apply_export_rules(exp_dicts, bullet_dicts, skill_dicts, edu_dicts, active_rules)
            experiences = transformed['experiences']
            bullets = transformed['bullets']
            skills = transformed['skills']
            education = transformed['education']
            section_order = transformed['section_order']

            # Extract personal header info from profile
            if profile_data.get('header_info'):
                header_info = profile_data['header_info']

    # Generate resume in requested format
    if export_format == 'txt':
        content_bytes, mimetype, filename = generate_resume_text(experiences, bullets, skills, education, timestamp, section_order, header_info)
    elif export_format == 'md':
        content_bytes, mimetype, filename = generate_resume_markdown(experiences, bullets, skills, education, timestamp, section_order, header_info)
    elif export_format == 'html':
        content_bytes, mimetype, filename = generate_resume_html(experiences, bullets, skills, education, timestamp, section_order, header_info)
    elif export_format == 'docx':
        content_bytes, mimetype, filename = generate_resume_docx(experiences, bullets, skills, education, timestamp, section_order, header_info)
    elif export_format == 'pdf':
        content_bytes, mimetype, filename = generate_resume_pdf(experiences, bullets, skills, education, timestamp, section_order, header_info)
    else:
        flash('Invalid export format', 'error')
        return redirect(url_for('export.export_select'))

    # Return as download
    return send_file(
        content_bytes,
        mimetype=mimetype,
        as_attachment=True,
        download_name=filename
    )


# ============================================================================
# RESUME FORMAT GENERATORS
# ============================================================================

def _prepare_export_data(experiences, bullets, skills, section_order=None, header_info=None):
    """Pre-process resume data shared across all export formats.

    Returns a dict with:
        section_order: list of section names
        header_info: header dict or None
        contact_parts: list of non-empty contact strings
        bullets_by_exp: {experience_id: [bullet, ...]}
        skills_by_cat: OrderedDict-style {category: [skill_name, ...]}
    """
    if section_order is None:
        section_order = ['experience', 'skills', 'education']

    contact_parts = []
    if header_info:
        for key in ('email', 'phone', 'location'):
            if header_info.get(key):
                contact_parts.append(header_info[key])

    bullets_by_exp = {}
    for bullet in bullets:
        exp_id = bullet['experience_id']
        if exp_id not in bullets_by_exp:
            bullets_by_exp[exp_id] = []
        bullets_by_exp[exp_id].append(bullet)

    skills_by_cat = {}
    for skill in skills:
        cat = skill['category'] or 'General'
        if cat not in skills_by_cat:
            skills_by_cat[cat] = []
        skills_by_cat[cat].append(skill['skill_name'])

    return {
        'section_order': section_order,
        'header_info': header_info,
        'contact_parts': contact_parts,
        'bullets_by_exp': bullets_by_exp,
        'skills_by_cat': skills_by_cat,
    }


def _render_sections(section_order, renderers):
    """Dispatch section renderers in order."""
    for section in section_order:
        renderer = renderers.get(section)
        if renderer:
            renderer()

def generate_resume_text(experiences, bullets, skills, education, timestamp, section_order=None, header_info=None):
    """Generate plain text resume from selected components"""
    data = _prepare_export_data(experiences, bullets, skills, section_order, header_info)
    output = []
    output.append("=" * 60)

    if header_info and header_info.get('name'):
        output.append(header_info['name'].upper())
        if data['contact_parts']:
            output.append(' | '.join(data['contact_parts']))
        if header_info.get('links'):
            output.append(header_info['links'])
    else:
        output.append("RESUME")

    output.append("=" * 60)
    output.append("")

    def _render_experience():
        if not experiences:
            return
        output.append("WORK EXPERIENCE")
        output.append("-" * 60)
        output.append("")

        for exp in experiences:
            output.append(f"{exp['job_title']}")
            output.append(f"{exp['company_name']}")
            output.append(f"{exp['start_date']} - {exp['end_date']}")
            if exp['location']:
                output.append(f"{exp['location']}")
            output.append("")

            if exp['description']:
                output.append(exp['description'])
                output.append("")

            if exp['id'] in data['bullets_by_exp']:
                for bullet in data['bullets_by_exp'][exp['id']]:
                    output.append(f"  • {bullet['bullet_text']}")
                output.append("")

        output.append("")

    def _render_skills():
        if not skills:
            return
        output.append("SKILLS")
        output.append("-" * 60)
        output.append("")

        for category, skill_list in data['skills_by_cat'].items():
            output.append(f"{category}: {', '.join(skill_list)}")

        output.append("")
        output.append("")

    def _render_education():
        if not education:
            return
        output.append("EDUCATION")
        output.append("-" * 60)
        output.append("")

        for edu in education:
            output.append(f"{edu['degree']} in {edu['field_of_study']}")
            output.append(f"{edu['school_name']}")
            output.append(f"{edu['graduation_year']}")
            if edu['location']:
                output.append(f"{edu['location']}")
            output.append("")

    _render_sections(data['section_order'], {
        'experience': _render_experience,
        'skills': _render_skills,
        'education': _render_education,
    })

    content = '\n'.join(output)
    content_bytes = BytesIO(content.encode('utf-8'))
    content_bytes.seek(0)

    return content_bytes, 'text/plain', f"resume_{timestamp}.txt"


def generate_resume_markdown(experiences, bullets, skills, education, timestamp, section_order=None, header_info=None):
    """Generate Markdown resume from selected components"""
    data = _prepare_export_data(experiences, bullets, skills, section_order, header_info)
    output = []

    if header_info and header_info.get('name'):
        output.append(f"# {header_info['name']}")
        if data['contact_parts']:
            output.append(' | '.join(data['contact_parts']))
        if header_info.get('links'):
            output.append(header_info['links'])
    else:
        output.append("# RESUME")

    output.append("")

    def _render_experience():
        if not experiences:
            return
        output.append("## WORK EXPERIENCE")
        output.append("")

        for exp in experiences:
            output.append(f"### {exp['job_title']}")
            output.append(f"**{exp['company_name']}**")
            output.append(f"*{exp['start_date']} - {exp['end_date']}*")
            if exp['location']:
                output.append(f"*{exp['location']}*")
            output.append("")

            if exp['description']:
                output.append(exp['description'])
                output.append("")

            if exp['id'] in data['bullets_by_exp']:
                for bullet in data['bullets_by_exp'][exp['id']]:
                    output.append(f"- {bullet['bullet_text']}")
                output.append("")

    def _render_skills():
        if not skills:
            return
        output.append("## SKILLS")
        output.append("")

        for category, skill_list in data['skills_by_cat'].items():
            output.append(f"**{category}:** {', '.join(skill_list)}")
        output.append("")

    def _render_education():
        if not education:
            return
        output.append("## EDUCATION")
        output.append("")

        for edu in education:
            output.append(f"### {edu['degree']} in {edu['field_of_study']}")
            output.append(f"**{edu['school_name']}**")
            output.append(f"*{edu['graduation_year']}*")
            if edu['location']:
                output.append(f"*{edu['location']}*")
            output.append("")

    _render_sections(data['section_order'], {
        'experience': _render_experience,
        'skills': _render_skills,
        'education': _render_education,
    })

    content = '\n'.join(output)
    content_bytes = BytesIO(content.encode('utf-8'))
    content_bytes.seek(0)

    return content_bytes, 'text/markdown', f"resume_{timestamp}.md"


def generate_resume_html(experiences, bullets, skills, education, timestamp, section_order=None, header_info=None):
    """Generate HTML resume from selected components"""
    data = _prepare_export_data(experiences, bullets, skills, section_order, header_info)
    esc = html_escape

    output = []
    output.append("<!DOCTYPE html>")
    output.append("<html lang='en'>")
    output.append("<head>")
    output.append("    <meta charset='UTF-8'>")
    output.append("    <meta name='viewport' content='width=device-width, initial-scale=1.0'>")
    output.append("    <title>Resume</title>")
    output.append("    <style>")
    output.append("        body { font-family: Arial, sans-serif; max-width: 800px; margin: 40px auto; padding: 20px; line-height: 1.6; }")
    output.append("        h1 { border-bottom: 3px solid #333; padding-bottom: 10px; }")
    output.append("        h2 { color: #333; border-bottom: 2px solid #666; padding-bottom: 5px; margin-top: 30px; }")
    output.append("        h3 { color: #555; margin-bottom: 5px; }")
    output.append("        .meta { color: #666; font-style: italic; margin: 5px 0; }")
    output.append("        ul { margin: 10px 0; }")
    output.append("        .header-contact { text-align: center; color: #555; margin: 5px 0; }")
    output.append("        .footer { text-align: center; margin-top: 40px; padding-top: 20px; border-top: 1px solid #ccc; color: #666; font-size: 0.9em; }")
    output.append("    </style>")
    output.append("</head>")
    output.append("<body>")

    if header_info and header_info.get('name'):
        output.append(f"    <h1 style='text-align: center;'>{esc(header_info['name'])}</h1>")
        if data['contact_parts']:
            escaped_contact = [esc(p) for p in data['contact_parts']]
            output.append(f"    <p class='header-contact'>{' | '.join(escaped_contact)}</p>")
        if header_info.get('links'):
            output.append(f"    <p class='header-contact'>{esc(header_info['links'])}</p>")
    else:
        output.append("    <h1>RESUME</h1>")

    def _render_experience():
        if not experiences:
            return
        output.append("    <h2>WORK EXPERIENCE</h2>")

        for exp in experiences:
            output.append(f"    <h3>{esc(exp['job_title'])}</h3>")
            output.append(f"    <div class='meta'><strong>{esc(exp['company_name'])}</strong></div>")
            output.append(f"    <div class='meta'>{esc(exp['start_date'])} - {esc(exp['end_date'])}")
            if exp['location']:
                output.append(f" | {esc(exp['location'])}")
            output.append("</div>")

            if exp['description']:
                output.append(f"    <p>{esc(exp['description'])}</p>")

            if exp['id'] in data['bullets_by_exp']:
                output.append("    <ul>")
                for bullet in data['bullets_by_exp'][exp['id']]:
                    output.append(f"        <li>{esc(bullet['bullet_text'])}</li>")
                output.append("    </ul>")

    def _render_skills():
        if not skills:
            return
        output.append("    <h2>SKILLS</h2>")

        for category, skill_list in data['skills_by_cat'].items():
            escaped_skills = ', '.join(esc(s) for s in skill_list)
            output.append(f"    <p><strong>{esc(category)}:</strong> {escaped_skills}</p>")

    def _render_education():
        if not education:
            return
        output.append("    <h2>EDUCATION</h2>")

        for edu in education:
            output.append(f"    <h3>{esc(edu['degree'])} in {esc(edu['field_of_study'])}</h3>")
            output.append(f"    <div class='meta'><strong>{esc(edu['school_name'])}</strong></div>")
            output.append(f"    <div class='meta'>{esc(edu['graduation_year'])}")
            if edu['location']:
                output.append(f" | {esc(edu['location'])}")
            output.append("</div>")

    _render_sections(data['section_order'], {
        'experience': _render_experience,
        'skills': _render_skills,
        'education': _render_education,
    })

    output.append("</body>")
    output.append("</html>")

    content = '\n'.join(output)
    content_bytes = BytesIO(content.encode('utf-8'))
    content_bytes.seek(0)

    return content_bytes, 'text/html', f"resume_{timestamp}.html"


def generate_resume_docx(experiences, bullets, skills, education, timestamp, section_order=None, header_info=None):
    """Generate DOCX resume from selected components"""
    data = _prepare_export_data(experiences, bullets, skills, section_order, header_info)
    doc = Document()

    if header_info and header_info.get('name'):
        title = doc.add_heading(header_info['name'], 0)
        title.alignment = 1  # Center

        if data['contact_parts']:
            p = doc.add_paragraph()
            p.alignment = 1
            p.add_run(' | '.join(data['contact_parts']))

        if header_info.get('links'):
            p = doc.add_paragraph()
            p.alignment = 1
            p.add_run(header_info['links'])
    else:
        title = doc.add_heading('RESUME', 0)
        title.alignment = 1  # Center alignment

    def _render_experience():
        if not experiences:
            return
        doc.add_heading('WORK EXPERIENCE', level=1)

        for exp in experiences:
            doc.add_heading(exp['job_title'], level=2)

            p = doc.add_paragraph()
            run = p.add_run(exp['company_name'])
            run.bold = True

            p = doc.add_paragraph()
            run = p.add_run(f"{exp['start_date']} - {exp['end_date']}")
            run.italic = True
            if exp['location']:
                run2 = p.add_run(f" | {exp['location']}")
                run2.italic = True

            if exp['description']:
                doc.add_paragraph(exp['description'])

            if exp['id'] in data['bullets_by_exp']:
                for bullet in data['bullets_by_exp'][exp['id']]:
                    doc.add_paragraph(bullet['bullet_text'], style='List Bullet')

    def _render_skills():
        if not skills:
            return
        doc.add_heading('SKILLS', level=1)

        for category, skill_list in data['skills_by_cat'].items():
            p = doc.add_paragraph()
            run = p.add_run(f"{category}: ")
            run.bold = True
            p.add_run(', '.join(skill_list))

    def _render_education():
        if not education:
            return
        doc.add_heading('EDUCATION', level=1)

        for edu in education:
            doc.add_heading(f"{edu['degree']} in {edu['field_of_study']}", level=2)

            p = doc.add_paragraph()
            run = p.add_run(edu['school_name'])
            run.bold = True

            p = doc.add_paragraph()
            run = p.add_run(edu['graduation_year'])
            run.italic = True
            if edu['location']:
                run2 = p.add_run(f" | {edu['location']}")
                run2.italic = True

    _render_sections(data['section_order'], {
        'experience': _render_experience,
        'skills': _render_skills,
        'education': _render_education,
    })

    # Save to BytesIO
    content_bytes = BytesIO()
    doc.save(content_bytes)
    content_bytes.seek(0)

    return content_bytes, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', f"resume_{timestamp}.docx"


def generate_resume_pdf(experiences, bullets, skills, education, timestamp, section_order=None, header_info=None):
    """Generate PDF resume from selected components"""
    data = _prepare_export_data(experiences, bullets, skills, section_order, header_info)
    esc = html_escape

    content_bytes = BytesIO()
    doc = SimpleDocTemplate(content_bytes, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        alignment=1,  # Center
        spaceAfter=6
    )

    contact_style = ParagraphStyle(
        'ContactInfo',
        parent=styles['Normal'],
        fontSize=10,
        alignment=1,  # Center
        spaceAfter=4
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=12
    )

    subheading_style = ParagraphStyle(
        'CustomSubHeading',
        parent=styles['Heading3'],
        fontSize=12,
        spaceAfter=6
    )

    # Title / Header
    if header_info and header_info.get('name'):
        story.append(Paragraph(esc(header_info['name']), title_style))
        if data['contact_parts']:
            escaped_contact = [esc(p) for p in data['contact_parts']]
            story.append(Paragraph(' | '.join(escaped_contact), contact_style))
        if header_info.get('links'):
            story.append(Paragraph(esc(header_info['links']), contact_style))
    else:
        story.append(Paragraph("RESUME", title_style))

    story.append(Spacer(1, 0.2*inch))

    def _render_experience():
        if not experiences:
            return
        story.append(Paragraph("WORK EXPERIENCE", heading_style))

        for exp in experiences:
            story.append(Paragraph(esc(exp['job_title']), subheading_style))
            story.append(Paragraph(f"<b>{esc(exp['company_name'])}</b>", styles['Normal']))

            location_date = f"<i>{esc(exp['start_date'])} - {esc(exp['end_date'])}"
            if exp['location']:
                location_date += f" | {esc(exp['location'])}"
            location_date += "</i>"
            story.append(Paragraph(location_date, styles['Normal']))

            if exp['description']:
                story.append(Spacer(1, 0.1*inch))
                story.append(Paragraph(esc(exp['description']), styles['Normal']))

            if exp['id'] in data['bullets_by_exp']:
                story.append(Spacer(1, 0.1*inch))
                for bullet in data['bullets_by_exp'][exp['id']]:
                    story.append(Paragraph(f"&bull; {esc(bullet['bullet_text'])}", styles['Normal']))

            story.append(Spacer(1, 0.2*inch))

    def _render_skills():
        if not skills:
            return
        story.append(Paragraph("SKILLS", heading_style))

        for category, skill_list in data['skills_by_cat'].items():
            escaped_skills = ', '.join(esc(s) for s in skill_list)
            story.append(Paragraph(f"<b>{esc(category)}:</b> {escaped_skills}", styles['Normal']))

        story.append(Spacer(1, 0.2*inch))

    def _render_education():
        if not education:
            return
        story.append(Paragraph("EDUCATION", heading_style))

        for edu in education:
            story.append(Paragraph(f"{esc(edu['degree'])} in {esc(edu['field_of_study'])}", subheading_style))
            story.append(Paragraph(f"<b>{esc(edu['school_name'])}</b>", styles['Normal']))

            grad_location = f"<i>{esc(edu['graduation_year'])}"
            if edu['location']:
                grad_location += f" | {esc(edu['location'])}"
            grad_location += "</i>"
            story.append(Paragraph(grad_location, styles['Normal']))
            story.append(Spacer(1, 0.1*inch))

    _render_sections(data['section_order'], {
        'experience': _render_experience,
        'skills': _render_skills,
        'education': _render_education,
    })

    # Build PDF
    doc.build(story)
    content_bytes.seek(0)

    return content_bytes, 'application/pdf', f"resume_{timestamp}.pdf"
